"""Document loader module for PDF, image, text, CSV, and DOCX files with OCR support"""
import os
import csv
import hashlib
from datetime import datetime, timezone
from typing import List, Dict, Optional, Tuple
from abc import ABC, abstractmethod

# text_cleaner disabled -- was stripping useful content during ingestion
# from src.text_cleaner import remove_boilerplate, clean_text


def _resolve_tesseract_path(path: str) -> str:
    """Resolve a Tesseract path that may be a directory or executable to the actual executable path."""
    if not path:
        return path
    if os.path.isdir(path):
        exe_path = os.path.join(path, "tesseract.exe")
        if os.path.isfile(exe_path):
            return exe_path
    return path

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

try:
    import docx as python_docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class DocumentLoader(ABC):
    """Abstract base class for document loaders"""
    
    @abstractmethod
    def load(self, file_path: str) -> List[Dict]:
        """
        Load a document and return extracted text with metadata.
        
        Returns:
            List of dicts with 'content', 'metadata' keys
        """
        pass
    
    @staticmethod
    def generate_doc_id(file_path: str) -> str:
        """Generate a unique document ID based on file path"""
        return hashlib.md5(file_path.encode()).hexdigest()[:12]

    @staticmethod
    def get_file_dates(file_path: str) -> Dict[str, str]:
        """Return file modification time and current ingestion timestamp as ISO strings."""
        file_mtime = ""
        try:
            mtime = os.path.getmtime(file_path)
            file_mtime = datetime.fromtimestamp(mtime, tz=timezone.utc).isoformat()
        except OSError:
            pass
        return {
            "file_modified_at": file_mtime,
            "ingested_at": datetime.now(tz=timezone.utc).isoformat(),
        }


class PDFLoader(DocumentLoader):
    """
    Loader for PDF files.
    Uses PyMuPDF for text extraction, falls back to OCR for scanned pages.
    """
    
    def __init__(
        self,
        min_text_length: int = 100,
        ocr_language: str = "eng+chi_sim+chi_tra",
        tesseract_path: Optional[str] = None
    ):
        """
        Initialize PDF loader.
        
        Args:
            min_text_length: Minimum text length before triggering OCR
            ocr_language: Language for OCR (default: English + Chinese Simplified + Traditional)
            tesseract_path: Path to Tesseract executable (Windows)
        """
        if not PYMUPDF_AVAILABLE:
            raise ImportError("PyMuPDF (fitz) is required. Install with: pip install pymupdf")
        
        self.min_text_length = min_text_length
        self.ocr_language = ocr_language
        
        if tesseract_path and TESSERACT_AVAILABLE:
            pytesseract.pytesseract.tesseract_cmd = _resolve_tesseract_path(tesseract_path)
    
    def load(self, file_path: str) -> List[Dict]:
        """
        Load a PDF file and extract text from all pages.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of documents (one per page) with content and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        
        documents = []
        doc_id = self.generate_doc_id(file_path)
        date_meta = self.get_file_dates(file_path)
        
        try:
            pdf_doc = fitz.open(file_path)
            total_pages = len(pdf_doc)
            raw_texts = []
            extraction_methods = []
            
            print(f"📄 Loading PDF: {os.path.basename(file_path)} ({total_pages} pages)")
            
            for page_num in range(total_pages):
                page = pdf_doc[page_num]
                text = page.get_text().strip()
                extraction_method = "text"
                
                if len(text) < self.min_text_length:
                    ocr_text = self._ocr_page(page)
                    if ocr_text and len(ocr_text) > len(text):
                        text = ocr_text
                        extraction_method = "ocr"
                
                raw_texts.append(text)
                extraction_methods.append(extraction_method)
            
            pdf_doc.close()

            # Remove repeated headers / footers across pages
            cleaned_texts = remove_boilerplate(raw_texts)

            for page_num, (text, method) in enumerate(
                zip(cleaned_texts, extraction_methods)
            ):
                text = clean_text(text)
                if text:
                    documents.append({
                        'content': text,
                        'metadata': {
                            'source': file_path,
                            'type': 'pdf',
                            'page': page_num + 1,
                            'total_pages': total_pages,
                            'extraction_method': method,
                            'parent_doc_id': doc_id,
                            **date_meta,
                        }
                    })
                    print(f"  📃 Page {page_num + 1}: {len(text)} chars ({method})")
            
        except Exception as e:
            raise Exception(f"Error loading PDF {file_path}: {e}")
        
        return documents
    
    def _ocr_page(self, page) -> Optional[str]:
        """
        Perform OCR on a PDF page.
        
        Args:
            page: PyMuPDF page object
            
        Returns:
            Extracted text or None if OCR is not available
        """
        if not TESSERACT_AVAILABLE:
            print("    ⚠️ Tesseract not available for OCR")
            return None
        
        try:
            # Render page to image at higher resolution for better OCR
            mat = fitz.Matrix(2, 2)  # 2x zoom for better quality
            pix = page.get_pixmap(matrix=mat)
            
            # Convert to PIL Image
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            
            # Perform OCR
            text = pytesseract.image_to_string(img, lang=self.ocr_language)
            return text.strip()
            
        except Exception as e:
            print(f"    ⚠️ OCR failed: {e}")
            return None


class ImageLoader(DocumentLoader):
    """
    Loader for image files using OCR.
    Supports PNG, JPG, JPEG, TIFF, BMP formats.
    """
    
    SUPPORTED_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp'}
    
    def __init__(
        self,
        ocr_language: str = "eng+chi_sim+chi_tra",
        tesseract_path: Optional[str] = None
    ):
        """
        Initialize image loader.
        
        Args:
            ocr_language: Language for OCR (default: English + Chinese Simplified + Traditional)
            tesseract_path: Path to Tesseract executable (Windows)
        """
        if not TESSERACT_AVAILABLE:
            raise ImportError(
                "pytesseract and Pillow are required. Install with: "
                "pip install pytesseract Pillow"
            )
        
        self.ocr_language = ocr_language
        
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = _resolve_tesseract_path(tesseract_path)
    
    def load(self, file_path: str) -> List[Dict]:
        """
        Load an image file and extract text using OCR.
        
        Args:
            file_path: Path to the image file
            
        Returns:
            List with single document containing extracted text and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Image file not found: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported image format: {ext}. "
                f"Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        doc_id = self.generate_doc_id(file_path)
        date_meta = self.get_file_dates(file_path)
        
        print(f"🖼️ Loading image: {os.path.basename(file_path)}")
        
        try:
            img = Image.open(file_path)
            
            if img.mode not in ('RGB', 'L'):
                img = img.convert('RGB')
            
            text = pytesseract.image_to_string(img, lang=self.ocr_language)
            text = text.strip()
            
            if not text:
                print(f"  ⚠️ No text extracted from image")
                return []
            
            print(f"  ✅ Extracted {len(text)} characters")
            
            return [{
                'content': text,
                'metadata': {
                    'source': file_path,
                    'type': 'image',
                    'format': ext[1:],
                    'extraction_method': 'ocr',
                    'parent_doc_id': doc_id,
                    'image_size': f"{img.width}x{img.height}",
                    **date_meta,
                }
            }]
            
        except Exception as e:
            raise Exception(f"Error loading image {file_path}: {e}")


class TextFileLoader(DocumentLoader):
    """
    Loader for plain text files (.txt).
    Reads the file content directly.
    """
    
    SUPPORTED_EXTENSIONS = {'.txt'}
    
    def load(self, file_path: str) -> List[Dict]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Text file not found: {file_path}")
        
        doc_id = self.generate_doc_id(file_path)
        date_meta = self.get_file_dates(file_path)
        
        print(f"📝 Loading text file: {os.path.basename(file_path)}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read().strip()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read().strip()
        
        if not text:
            print(f"  ⚠️ Text file is empty")
            return []
        
        print(f"  ✅ Read {len(text)} characters")
        
        return [{
            'content': text,
            'metadata': {
                'source': file_path,
                'type': 'text',
                'format': 'txt',
                'extraction_method': 'direct_read',
                'parent_doc_id': doc_id,
                **date_meta,
            }
        }]


class CSVLoader(DocumentLoader):
    """
    Loader for CSV files (.csv).
    Reads CSV rows and converts them to readable text.
    """
    
    SUPPORTED_EXTENSIONS = {'.csv'}
    
    def load(self, file_path: str) -> List[Dict]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"CSV file not found: {file_path}")
        
        doc_id = self.generate_doc_id(file_path)
        date_meta = self.get_file_dates(file_path)
        
        print(f"📊 Loading CSV file: {os.path.basename(file_path)}")
        
        try:
            rows = []
            with open(file_path, 'r', encoding='utf-8', newline='') as f:
                reader = csv.reader(f)
                headers = next(reader, None)
                
                if headers is None:
                    print(f"  ⚠️ CSV file is empty")
                    return []
                
                row_count = 0
                for row in reader:
                    if row:
                        row_text = ", ".join(
                            f"{headers[i]}: {row[i]}" 
                            for i in range(min(len(headers), len(row)))
                            if row[i].strip()
                        )
                        if row_text:
                            rows.append(row_text)
                            row_count += 1
            
            if not rows:
                print(f"  ⚠️ No data rows found in CSV")
                return []
            
            text = f"Headers: {', '.join(headers)}\n\n" + "\n".join(rows)
            print(f"  ✅ Read {row_count} rows, {len(text)} characters")
            
            return [{
                'content': text,
                'metadata': {
                    'source': file_path,
                    'type': 'csv',
                    'format': 'csv',
                    'extraction_method': 'csv_reader',
                    'parent_doc_id': doc_id,
                    'row_count': row_count,
                    'columns': headers,
                    **date_meta,
                }
            }]
            
        except Exception as e:
            raise Exception(f"Error loading CSV {file_path}: {e}")


class DocxLoader(DocumentLoader):
    """
    Loader for DOCX files (.docx).
    Uses python-docx library to extract text from Word documents.
    """
    
    SUPPORTED_EXTENSIONS = {'.docx'}
    
    def __init__(self):
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for .docx support. "
                "Install with: pip install python-docx"
            )
    
    def load(self, file_path: str) -> List[Dict]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        doc_id = self.generate_doc_id(file_path)
        date_meta = self.get_file_dates(file_path)
        
        print(f"📄 Loading DOCX file: {os.path.basename(file_path)}")
        
        try:
            doc = python_docx.Document(file_path)
            
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            if not paragraphs:
                print(f"  ⚠️ No text extracted from DOCX")
                return []
            
            text = "\n\n".join(paragraphs)
            print(f"  ✅ Extracted {len(text)} characters from {len(paragraphs)} paragraphs")
            
            return [{
                'content': text,
                'metadata': {
                    'source': file_path,
                    'type': 'docx',
                    'format': 'docx',
                    'extraction_method': 'python_docx',
                    'parent_doc_id': doc_id,
                    'paragraph_count': len(paragraphs),
                    **date_meta,
                }
            }]
            
        except Exception as e:
            raise Exception(f"Error loading DOCX {file_path}: {e}")


class DocumentLoaderFactory:
    """
    Factory class that auto-detects file type and returns appropriate loader.
    """
    
    PDF_EXTENSIONS = {'.pdf'}
    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp'}
    TEXT_EXTENSIONS = {'.txt'}
    CSV_EXTENSIONS = {'.csv'}
    DOCX_EXTENSIONS = {'.docx'}
    
    def __init__(
        self,
        min_text_length: int = 100,
        ocr_language: str = "eng+chi_sim+chi_tra",
        tesseract_path: Optional[str] = None
    ):
        """
        Initialize the factory with common configuration.
        
        Args:
            min_text_length: Minimum text length before triggering OCR for PDFs
            ocr_language: Language for OCR (default: English + Chinese Simplified + Traditional)
            tesseract_path: Path to Tesseract executable
        """
        self.min_text_length = min_text_length
        self.ocr_language = ocr_language
        self.tesseract_path = tesseract_path
    
    def get_loader(self, file_path: str) -> DocumentLoader:
        """
        Get the appropriate loader for a file based on its extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Appropriate DocumentLoader instance
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext in self.PDF_EXTENSIONS:
            return PDFLoader(
                min_text_length=self.min_text_length,
                ocr_language=self.ocr_language,
                tesseract_path=self.tesseract_path
            )
        elif ext in self.IMAGE_EXTENSIONS:
            return ImageLoader(
                ocr_language=self.ocr_language,
                tesseract_path=self.tesseract_path
            )
        elif ext in self.TEXT_EXTENSIONS:
            return TextFileLoader()
        elif ext in self.CSV_EXTENSIONS:
            return CSVLoader()
        elif ext in self.DOCX_EXTENSIONS:
            return DocxLoader()
        else:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: PDF ({', '.join(self.PDF_EXTENSIONS)}), "
                f"Images ({', '.join(self.IMAGE_EXTENSIONS)}), "
                f"Text ({', '.join(self.TEXT_EXTENSIONS)}), "
                f"CSV ({', '.join(self.CSV_EXTENSIONS)}), "
                f"DOCX ({', '.join(self.DOCX_EXTENSIONS)})"
            )
    
    def load(self, file_path: str) -> List[Dict]:
        """
        Load a document using the appropriate loader.
        
        Args:
            file_path: Path to the file
            
        Returns:
            List of documents with content and metadata
        """
        loader = self.get_loader(file_path)
        return loader.load(file_path)
    
    def load_directory(
        self,
        directory: str,
        extensions: Optional[List[str]] = None,
        recursive: bool = False
    ) -> List[Dict]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory: Path to the directory
            extensions: Optional list of extensions to filter (e.g., ['.pdf', '.png'])
            recursive: Whether to search subdirectories
            
        Returns:
            List of all documents from all files
        """
        if not os.path.isdir(directory):
            raise NotADirectoryError(f"Directory not found: {directory}")
        
        # Default to all supported extensions
        if extensions is None:
            extensions = list(self.PDF_EXTENSIONS | self.IMAGE_EXTENSIONS | self.TEXT_EXTENSIONS | self.CSV_EXTENSIONS | self.DOCX_EXTENSIONS)
        
        # Normalize extensions
        extensions = [ext.lower() if ext.startswith('.') else f'.{ext.lower()}' 
                      for ext in extensions]
        
        all_documents = []
        files_processed = 0
        
        print(f"📁 Scanning directory: {directory}")
        print(f"   Extensions: {', '.join(extensions)}")
        
        if recursive:
            for root, _, files in os.walk(directory):
                for filename in files:
                    file_path = os.path.join(root, filename)
                    ext = os.path.splitext(filename)[1].lower()
                    if ext in extensions:
                        try:
                            docs = self.load(file_path)
                            all_documents.extend(docs)
                            files_processed += 1
                        except Exception as e:
                            print(f"   ⚠️ Error loading {filename}: {e}")
        else:
            for filename in os.listdir(directory):
                file_path = os.path.join(directory, filename)
                if os.path.isfile(file_path):
                    ext = os.path.splitext(filename)[1].lower()
                    if ext in extensions:
                        try:
                            docs = self.load(file_path)
                            all_documents.extend(docs)
                            files_processed += 1
                        except Exception as e:
                            print(f"   ⚠️ Error loading {filename}: {e}")
        
        print(f"✅ Processed {files_processed} files, extracted {len(all_documents)} document segments")
        
        return all_documents


def check_dependencies() -> Dict[str, bool]:
    """
    Check which dependencies are available.
    
    Returns:
        Dict with dependency availability status
    """
    status = {
        'pymupdf': PYMUPDF_AVAILABLE,
        'pytesseract': TESSERACT_AVAILABLE,
        'pdf2image': PDF2IMAGE_AVAILABLE,
        'python_docx': PYTHON_DOCX_AVAILABLE,
    }
    
    # Check if Tesseract executable is accessible
    if TESSERACT_AVAILABLE:
        # Try to configure tesseract path from environment/config before checking
        _configured_path = os.environ.get("TESSERACT_PATH", "")
        if _configured_path and pytesseract.pytesseract.tesseract_cmd == "tesseract":
            pytesseract.pytesseract.tesseract_cmd = _resolve_tesseract_path(_configured_path)
        try:
            pytesseract.get_tesseract_version()
            status['tesseract_executable'] = True
        except Exception:
            status['tesseract_executable'] = False
    else:
        status['tesseract_executable'] = False
    
    return status


def print_dependency_status():
    """Print the status of all dependencies"""
    status = check_dependencies()
    
    print("\n📦 Document Loader Dependencies:")
    print("-" * 40)
    
    for dep, available in status.items():
        icon = "✅" if available else "❌"
        print(f"  {icon} {dep}: {'Available' if available else 'Not available'}")
    
    if not status['tesseract_executable']:
        print("\n⚠️  Tesseract OCR is not installed or not in PATH.")
        print("   Windows: Download from https://github.com/UB-Mannheim/tesseract/wiki")
        print("   Then add to PATH or set TESSERACT_PATH in config.py")
    
    print("-" * 40)
